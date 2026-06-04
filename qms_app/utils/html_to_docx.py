from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup

def convert_html_to_docx(html, output_file):

    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    for element in soup.find_all(["p", "h1", "h2", "h3", "h4", "table"]):

        if element.name == "p":
            p = doc.add_paragraph()
            add_html_text_formatting(p, element)

        elif element.name.startswith("h"):
            level = int(element.name[1])
            p = doc.add_heading(level=level)
            add_html_text_formatting(p, element)

        elif element.name == "table":

            rows = element.find_all("tr")
            if not rows:
                continue

            cols = max(len(r.find_all(["td", "th"])) for r in rows)

            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"

            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])

                for j, cell in enumerate(cells):
                    doc_cell = table.rows[i].cells[j]
                    doc_cell.text = ""
                    p = doc_cell.paragraphs[0]
                    add_html_text_formatting(p, cell)

    doc.save(output_file)

def add_html_text_formatting(paragraph, html_element):

    for child in html_element.children:

        if child.name is None:
            paragraph.add_run(str(child).strip())

        else:
            text = child.get_text(strip=True)
            if not text:
                continue

            run = paragraph.add_run(text)

            style = child.get("style", "")

            if child.name in ["b", "strong"] or "font-weight:bold" in style:
                run.bold = True

            if child.name in ["i", "em"] or "font-style:italic" in style:
                run.italic = True

            if child.name == "u" or "text-decoration:underline" in style:
                run.underline = True

            if "font-size" in style:
                try:
                    size = style.split("font-size:")[1].split(";")[0]
                    size = size.replace("px", "").strip()
                    run.font.size = Pt(int(size) * 0.75)
                except:
                    pass